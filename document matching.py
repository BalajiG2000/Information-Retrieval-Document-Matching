## 181B069 BALAJI G
from docx import Document
import pandas as pd
document = Document()
document = Document(r'C:\Users\Admin\Downloads\Assignment task 1.docx')
table=document.tables[0]
data = [[cell.text for cell in row.cells] for row in table.rows]
df = pd.DataFrame(data)
print(df)
query=input()
a=[]
for col in df:
    a.append(df[col].str.contains(query))
for ind in df.index:
    if a[1][ind] == True:
        print(df[0][ind])        
